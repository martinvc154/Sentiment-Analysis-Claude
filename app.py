import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.figure_factory as ff
from scipy.cluster import hierarchy
from scipy.spatial.distance import pdist
from scipy import stats
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sns
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import tempfile
import shutil

# Set page configuration
st.set_page_config(
    page_title="Sentiment Analysis Dashboard",
    page_icon="📊",
    layout="wide"
)

# Initialize session state
if 'data' not in st.session_state:
    st.session_state.data = None
if 'org_column' not in st.session_state:
    st.session_state.org_column = None
if 'question_column' not in st.session_state:
    st.session_state.question_column = None
if 'sentiment_column' not in st.session_state:
    st.session_state.sentiment_column = None

# Create exports directory if it doesn't exist
if not os.path.exists('exports'):
    os.makedirs('exports')

# Title
st.title("📊 Sentiment Analysis Dashboard")

# Create tabs
tab1, tab2, tab3, tab4 = st.tabs(["📁 Data Upload", "🔗 Column Mapping", "📈 Visualizations", "📊 Export"])

# TAB 1: Data Upload
with tab1:
    st.header("Upload Your Data")
    st.markdown("Upload a CSV or Excel file containing your sentiment analysis data.")

    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['csv', 'xlsx', 'xls'],
        help="Upload a CSV or Excel file with your sentiment data"
    )

    if uploaded_file is not None:
        try:
            # Read the file
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)

            st.session_state.data = df
            st.success(f"✅ File uploaded successfully! {len(df)} rows loaded.")

            # Show data preview
            st.subheader("Data Preview")
            st.dataframe(df.head(10), use_container_width=True)

            # Show basic statistics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total Rows", len(df))
            with col2:
                st.metric("Total Columns", len(df.columns))
            with col3:
                st.metric("Memory Usage", f"{df.memory_usage(deep=True).sum() / 1024:.1f} KB")

        except Exception as e:
            st.error(f"❌ Error loading file: {str(e)}")

    elif st.session_state.data is not None:
        st.info("✅ Data already loaded. Upload a new file to replace it.")
        st.dataframe(st.session_state.data.head(10), use_container_width=True)

# TAB 2: Column Mapping
with tab2:
    st.header("Map Your Columns")
    st.markdown("Select which columns in your dataset correspond to organizations, questions, and sentiment scores.")

    if st.session_state.data is None:
        st.warning("⚠️ Please upload data in the 'Data Upload' tab first.")
    else:
        df = st.session_state.data
        columns = df.columns.tolist()

        col1, col2, col3 = st.columns(3)

        with col1:
            st.subheader("Organization Column")
            org_col = st.selectbox(
                "Select the column containing organization names:",
                options=['None'] + columns,
                index=0 if st.session_state.org_column is None else columns.index(st.session_state.org_column) + 1 if st.session_state.org_column in columns else 0,
                key='org_select'
            )
            if org_col != 'None':
                st.session_state.org_column = org_col
                st.info(f"Unique organizations: {df[org_col].nunique()}")
            else:
                st.session_state.org_column = None

        with col2:
            st.subheader("Question Column")
            q_col = st.selectbox(
                "Select the column containing questions:",
                options=['None'] + columns,
                index=0 if st.session_state.question_column is None else columns.index(st.session_state.question_column) + 1 if st.session_state.question_column in columns else 0,
                key='question_select'
            )
            if q_col != 'None':
                st.session_state.question_column = q_col
                st.info(f"Unique questions: {df[q_col].nunique()}")
            else:
                st.session_state.question_column = None

        with col3:
            st.subheader("Sentiment Column")
            sent_col = st.selectbox(
                "Select the column containing sentiment scores:",
                options=['None'] + columns,
                index=0 if st.session_state.sentiment_column is None else columns.index(st.session_state.sentiment_column) + 1 if st.session_state.sentiment_column in columns else 0,
                key='sentiment_select'
            )
            if sent_col != 'None':
                st.session_state.sentiment_column = sent_col
                # Try to show statistics if numeric
                try:
                    st.info(f"Range: [{df[sent_col].min():.2f}, {df[sent_col].max():.2f}]")
                except:
                    st.info("Non-numeric values detected")
            else:
                st.session_state.sentiment_column = None

        # Show mapping summary
        st.divider()
        st.subheader("Mapping Summary")
        mapping_df = pd.DataFrame({
            'Field': ['Organization', 'Question', 'Sentiment'],
            'Mapped Column': [
                st.session_state.org_column or '❌ Not mapped',
                st.session_state.question_column or '❌ Not mapped',
                st.session_state.sentiment_column or '❌ Not mapped'
            ]
        })
        st.dataframe(mapping_df, use_container_width=True, hide_index=True)

# TAB 3: Visualizations
with tab3:
    st.header("Visualizations")

    if st.session_state.data is None:
        st.warning("⚠️ Please upload data in the 'Data Upload' tab first.")
    else:
        # Chart selector
        st.subheader("Select Visualization Type")
        chart_type = st.radio(
            "Choose a visualization:",
            options=["Heatmap", "Divergence Plot", "Emotion Profiles", "Clusters"],
            horizontal=True,
            help="Select the type of visualization you want to create"
        )

        st.divider()

        # HEATMAP VISUALIZATION
        if chart_type == "Heatmap":
            # Check if required columns are mapped
            if st.session_state.org_column is None or st.session_state.question_column is None:
                st.error("❌ Both Organization and Question columns must be mapped to display the heatmap.")
                st.info("👈 Please go to the 'Column Mapping' tab and map both columns.")
            elif st.session_state.sentiment_column is None:
                st.error("❌ Sentiment column must be mapped to display the heatmap.")
                st.info("👈 Please go to the 'Column Mapping' tab and map the sentiment column.")
            else:
                df = st.session_state.data
                org_col = st.session_state.org_column
                q_col = st.session_state.question_column
                sent_col = st.session_state.sentiment_column

                try:
                    # Create pivot table with average sentiment scores
                    pivot_data = df.pivot_table(
                        values=sent_col,
                        index=org_col,
                        columns=q_col,
                        aggfunc='mean'
                    )

                    if pivot_data.empty:
                        st.error("❌ No data available for heatmap. Please check your data and column mappings.")
                    else:
                        # Replace NaN with 0 for clustering
                        pivot_filled = pivot_data.fillna(0)

                        # Perform hierarchical clustering
                        if len(pivot_filled) > 1 and len(pivot_filled.columns) > 1:
                            # Cluster rows (organizations)
                            row_linkage = hierarchy.linkage(
                                pdist(pivot_filled, metric='euclidean'),
                                method='ward'
                            )
                            row_dendro = hierarchy.dendrogram(row_linkage, no_plot=True)
                            row_order = row_dendro['leaves']

                            # Cluster columns (questions)
                            col_linkage = hierarchy.linkage(
                                pdist(pivot_filled.T, metric='euclidean'),
                                method='ward'
                            )
                            col_dendro = hierarchy.dendrogram(col_linkage, no_plot=True)
                            col_order = col_dendro['leaves']

                            # Reorder the data based on clustering
                            pivot_clustered = pivot_filled.iloc[row_order, col_order]
                        else:
                            pivot_clustered = pivot_filled
                            row_dendro = None
                            col_dendro = None

                        # Create the heatmap using Plotly
                        st.subheader("Sentiment Heatmap with Hierarchical Clustering")

                        # Create dendrogram figures if clustering was performed
                        if row_dendro is not None and col_dendro is not None:
                            # Create figure with dendrograms
                            fig = ff.create_dendrogram(
                                pivot_filled.values,
                                orientation='left',
                                labels=pivot_filled.index.tolist()
                            )
                            fig.update_layout(width=200, height=600)

                            # For simplicity, show heatmap without integrated dendrograms
                            # (Plotly's create_dendrogram doesn't easily integrate with heatmaps)
                            # Instead, create a standalone heatmap with clustered data

                            fig = go.Figure(data=go.Heatmap(
                                z=pivot_clustered.values,
                                x=pivot_clustered.columns.tolist(),
                                y=pivot_clustered.index.tolist(),
                                colorscale='RdYlGn',
                                zmid=0,
                                zmin=-1,
                                zmax=1,
                                colorbar=dict(title="Sentiment Score"),
                                text=np.round(pivot_clustered.values, 2),
                                texttemplate='%{text}',
                                textfont={"size": 10},
                                hovertemplate='Organization: %{y}<br>Question: %{x}<br>Sentiment: %{z:.2f}<extra></extra>'
                            ))

                            fig.update_layout(
                                title="Sentiment Heatmap (Hierarchically Clustered)",
                                xaxis_title="Questions",
                                yaxis_title="Organizations",
                                height=max(400, len(pivot_clustered) * 30),
                                width=max(800, len(pivot_clustered.columns) * 50),
                                xaxis={'side': 'bottom'},
                                yaxis={'side': 'left'}
                            )
                        else:
                            # Simple heatmap without clustering (not enough data points)
                            fig = go.Figure(data=go.Heatmap(
                                z=pivot_clustered.values,
                                x=pivot_clustered.columns.tolist(),
                                y=pivot_clustered.index.tolist(),
                                colorscale='RdYlGn',
                                zmid=0,
                                zmin=-1,
                                zmax=1,
                                colorbar=dict(title="Sentiment Score"),
                                text=np.round(pivot_clustered.values, 2),
                                texttemplate='%{text}',
                                textfont={"size": 10},
                                hovertemplate='Organization: %{y}<br>Question: %{x}<br>Sentiment: %{z:.2f}<extra></extra>'
                            ))

                            fig.update_layout(
                                title="Sentiment Heatmap",
                                xaxis_title="Questions",
                                yaxis_title="Organizations",
                                height=max(400, len(pivot_clustered) * 30),
                                width=max(800, len(pivot_clustered.columns) * 50)
                            )

                        st.plotly_chart(fig, use_container_width=True)

                        # Store the pivot data for export
                        st.session_state.heatmap_data = pivot_clustered
                        st.session_state.row_linkage = row_linkage if len(pivot_filled) > 1 else None
                        st.session_state.col_linkage = col_linkage if len(pivot_filled.columns) > 1 else None

                        # EXPORT SECTION
                        st.divider()
                        st.subheader("Export Options")

                        col1, col2, col3 = st.columns([2, 2, 2])

                        with col1:
                            dpi = st.selectbox(
                                "Select DPI:",
                                options=[150, 300, 600],
                                index=1,
                                help="Higher DPI = better quality but larger file size"
                            )

                        with col2:
                            size_options = {
                                "7x5 inches": (7, 5),
                                "10x7 inches": (10, 7),
                                "14x10 inches": (14, 10)
                            }
                            size_label = st.selectbox(
                                "Select Size:",
                                options=list(size_options.keys()),
                                index=1
                            )
                            fig_size = size_options[size_label]

                        with col3:
                            st.write("")  # Spacing
                            st.write("")  # Spacing

                        # Export buttons
                        col1, col2 = st.columns(2)

                        def create_publication_heatmap(data, row_linkage, col_linkage, figsize, dpi):
                            """Create a publication-quality heatmap with dendrograms using matplotlib"""

                            # Set font to Times New Roman
                            matplotlib.rcParams['font.family'] = 'serif'
                            matplotlib.rcParams['font.serif'] = ['Times New Roman']
                            matplotlib.rcParams['font.size'] = 12
                            matplotlib.rcParams['axes.linewidth'] = 2

                            # Create figure with subplots for dendrograms and heatmap
                            if row_linkage is not None and col_linkage is not None:
                                fig = plt.figure(figsize=figsize, dpi=dpi)

                                # Create grid spec
                                gs = fig.add_gridspec(2, 2, width_ratios=[0.2, 1], height_ratios=[0.2, 1],
                                                     hspace=0.05, wspace=0.05,
                                                     left=0.15, right=0.95, top=0.95, bottom=0.15)

                                # Top dendrogram (columns)
                                ax_top = fig.add_subplot(gs[0, 1])
                                dendro_top = hierarchy.dendrogram(
                                    col_linkage,
                                    ax=ax_top,
                                    color_threshold=0,
                                    above_threshold_color='black',
                                    no_labels=True
                                )
                                ax_top.set_xticks([])
                                ax_top.set_yticks([])
                                ax_top.spines['top'].set_visible(False)
                                ax_top.spines['right'].set_visible(False)
                                ax_top.spines['bottom'].set_visible(False)
                                ax_top.spines['left'].set_visible(False)

                                # Left dendrogram (rows)
                                ax_left = fig.add_subplot(gs[1, 0])
                                dendro_left = hierarchy.dendrogram(
                                    row_linkage,
                                    ax=ax_left,
                                    orientation='left',
                                    color_threshold=0,
                                    above_threshold_color='black',
                                    no_labels=True
                                )
                                ax_left.set_xticks([])
                                ax_left.set_yticks([])
                                ax_left.spines['top'].set_visible(False)
                                ax_left.spines['right'].set_visible(False)
                                ax_left.spines['bottom'].set_visible(False)
                                ax_left.spines['left'].set_visible(False)

                                # Main heatmap
                                ax_heatmap = fig.add_subplot(gs[1, 1])
                            else:
                                # Simple heatmap without dendrograms
                                fig, ax_heatmap = plt.subplots(figsize=figsize, dpi=dpi)

                            # Create heatmap
                            im = ax_heatmap.imshow(
                                data.values,
                                cmap='RdYlGn',
                                aspect='auto',
                                vmin=-1,
                                vmax=1,
                                interpolation='nearest'
                            )

                            # Set ticks and labels
                            ax_heatmap.set_xticks(np.arange(len(data.columns)))
                            ax_heatmap.set_yticks(np.arange(len(data.index)))
                            ax_heatmap.set_xticklabels(data.columns, rotation=45, ha='right', fontsize=10)
                            ax_heatmap.set_yticklabels(data.index, fontsize=10)

                            # Add gridlines
                            ax_heatmap.set_xticks(np.arange(len(data.columns)) - 0.5, minor=True)
                            ax_heatmap.set_yticks(np.arange(len(data.index)) - 0.5, minor=True)
                            ax_heatmap.grid(which='minor', color='white', linestyle='-', linewidth=2)

                            # Add colorbar
                            cbar = plt.colorbar(im, ax=ax_heatmap, fraction=0.046, pad=0.04)
                            cbar.set_label('Sentiment Score', rotation=270, labelpad=25, fontsize=12, weight='bold')
                            cbar.outline.set_linewidth(2)

                            # Add cell values
                            for i in range(len(data.index)):
                                for j in range(len(data.columns)):
                                    value = data.values[i, j]
                                    if not np.isnan(value):
                                        text_color = 'white' if abs(value) > 0.5 else 'black'
                                        ax_heatmap.text(j, i, f'{value:.2f}',
                                                       ha='center', va='center',
                                                       color=text_color, fontsize=8, weight='bold')

                            # Labels
                            ax_heatmap.set_xlabel('Questions', fontsize=14, weight='bold')
                            ax_heatmap.set_ylabel('Organizations', fontsize=14, weight='bold')

                            # Title
                            if row_linkage is not None and col_linkage is not None:
                                fig.suptitle('Sentiment Heatmap with Hierarchical Clustering',
                                           fontsize=16, weight='bold', y=0.98)
                            else:
                                ax_heatmap.set_title('Sentiment Heatmap', fontsize=16, weight='bold', pad=20)

                            # Adjust spines
                            for spine in ax_heatmap.spines.values():
                                spine.set_linewidth(2)

                            plt.tight_layout()
                            return fig

                        with col1:
                            if st.button("📄 Export as PDF", use_container_width=True):
                                with st.spinner("Generating PDF..."):
                                    try:
                                        fig = create_publication_heatmap(
                                            st.session_state.heatmap_data,
                                            st.session_state.row_linkage,
                                            st.session_state.col_linkage,
                                            fig_size,
                                            dpi
                                        )

                                        # Save to file
                                        output_path = 'exports/sentiment_heatmap.pdf'
                                        fig.savefig(output_path, format='pdf', bbox_inches='tight', dpi=dpi)
                                        plt.close(fig)

                                        # Provide download button
                                        with open(output_path, 'rb') as f:
                                            st.download_button(
                                                label="⬇️ Download PDF",
                                                data=f,
                                                file_name="sentiment_heatmap.pdf",
                                                mime="application/pdf",
                                                use_container_width=True
                                            )
                                        st.success(f"✅ PDF saved to {output_path}")
                                    except Exception as e:
                                        st.error(f"❌ Error creating PDF: {str(e)}")

                        with col2:
                            if st.button("🖼️ Export as PNG", use_container_width=True):
                                with st.spinner("Generating PNG..."):
                                    try:
                                        fig = create_publication_heatmap(
                                            st.session_state.heatmap_data,
                                            st.session_state.row_linkage,
                                            st.session_state.col_linkage,
                                            fig_size,
                                            dpi
                                        )

                                        # Save to file
                                        output_path = 'exports/sentiment_heatmap.png'
                                        fig.savefig(output_path, format='png', bbox_inches='tight', dpi=dpi)
                                        plt.close(fig)

                                        # Provide download button
                                        with open(output_path, 'rb') as f:
                                            st.download_button(
                                                label="⬇️ Download PNG",
                                                data=f,
                                                file_name="sentiment_heatmap.png",
                                                mime="image/png",
                                                use_container_width=True
                                            )
                                        st.success(f"✅ PNG saved to {output_path}")
                                    except Exception as e:
                                        st.error(f"❌ Error creating PNG: {str(e)}")

                except Exception as e:
                    st.error(f"❌ Error creating heatmap: {str(e)}")
                    st.info("Please check that your sentiment column contains numeric values.")

        # Placeholder messages for other visualizations
        elif chart_type == "Divergence Plot":
            st.info("📊 Divergence Plot visualization coming soon!")
            st.markdown("This will show sentiment divergence across different dimensions.")

        elif chart_type == "Emotion Profiles":
            st.info("😊 Emotion Profiles visualization coming soon!")
            st.markdown("This will display emotion distributions and profiles.")

        elif chart_type == "Clusters":
            st.info("🔍 Clusters visualization coming soon!")
            st.markdown("This will show clustering analysis of sentiment patterns.")

# TAB 4: Export
with tab4:
    st.header("Export Statistics and Data")

    if st.session_state.data is None:
        st.warning("⚠️ Please upload data in the 'Data Upload' tab first.")
    else:
        df = st.session_state.data

        # Helper function to calculate descriptive statistics with confidence intervals
        def calculate_descriptive_stats(df, columns):
            """Calculate descriptive statistics with 95% confidence intervals"""
            stats_data = []

            for col in columns:
                if col in df.columns:
                    data = df[col].dropna()
                    if len(data) > 0:
                        n = len(data)
                        mean = data.mean()
                        std = data.std()
                        min_val = data.min()
                        max_val = data.max()

                        # Calculate 95% confidence interval
                        if n > 1:
                            se = std / np.sqrt(n)
                            ci = stats.t.interval(0.95, n-1, loc=mean, scale=se)
                            ci_str = f"[{ci[0]:.3f}, {ci[1]:.3f}]"
                        else:
                            ci_str = "N/A"

                        stats_data.append({
                            'Variable': col,
                            'N': n,
                            'Mean': f"{mean:.3f}",
                            'SD': f"{std:.3f}",
                            'Min': f"{min_val:.3f}",
                            'Max': f"{max_val:.3f}",
                            '95% CI': ci_str
                        })

            return pd.DataFrame(stats_data)

        # Helper function to calculate organization summary
        def calculate_org_summary(df, org_col, sent_col):
            """Calculate organization-level summary statistics"""
            summary_data = []

            # Detect emotion columns
            emotion_cols = [col for col in df.columns if any(emotion in col.lower()
                           for emotion in ['joy', 'sadness', 'anger', 'fear', 'surprise', 'disgust'])]

            # Detect divergence column
            div_col = None
            for col in df.columns:
                if 'divergence' in col.lower():
                    div_col = col
                    break

            for org in df[org_col].unique():
                org_data = df[df[org_col] == org]
                n = len(org_data)

                # Sentiment statistics
                sent_mean = org_data[sent_col].mean()
                sent_std = org_data[sent_col].std()

                # Divergence statistics (if available)
                if div_col and div_col in df.columns:
                    div_mean = org_data[div_col].mean()
                    div_std = org_data[div_col].std()
                    div_str = f"{div_mean:.3f} ({div_std:.3f})"
                else:
                    div_str = "N/A"

                # Dominant emotion (if emotions available)
                if emotion_cols:
                    emotion_means = {col: org_data[col].mean() for col in emotion_cols if col in org_data.columns}
                    if emotion_means:
                        dominant = max(emotion_means, key=emotion_means.get)
                        # Clean up column name to get emotion name
                        dominant_emotion = dominant.replace('emotion_', '').replace('_', ' ').title()
                    else:
                        dominant_emotion = "N/A"
                else:
                    dominant_emotion = "N/A"

                summary_data.append({
                    'Organization': org,
                    'N': n,
                    'Sentiment M(SD)': f"{sent_mean:.3f} ({sent_std:.3f})",
                    'Divergence M(SD)': div_str,
                    'Dominant Emotion': dominant_emotion
                })

            return pd.DataFrame(summary_data)

        # Helper function to create LaTeX table
        def create_latex_table(df, caption, label):
            """Generate LaTeX table with booktabs style"""
            latex = "\\begin{table}[htbp]\n"
            latex += "\\centering\n"
            latex += f"\\caption{{{caption}}}\n"
            latex += f"\\label{{{label}}}\n"

            # Column alignment
            n_cols = len(df.columns)
            alignment = "l" + "c" * (n_cols - 1)
            latex += f"\\begin{{tabular}}{{{alignment}}}\n"
            latex += "\\toprule\n"

            # Header
            latex += " & ".join(df.columns) + " \\\\\n"
            latex += "\\midrule\n"

            # Rows
            for _, row in df.iterrows():
                latex += " & ".join(str(val) for val in row.values) + " \\\\\n"

            latex += "\\bottomrule\n"
            latex += "\\end{tabular}\n"
            latex += "\\end{table}\n"

            return latex

        # Helper function to create Word document with tables
        def create_word_document(desc_stats_df, org_summary_df=None):
            """Create a Word document with formatted tables"""
            doc = Document()

            # Title
            title = doc.add_heading('Sentiment Analysis Statistics', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section 1: Descriptive Statistics
            doc.add_heading('Descriptive Statistics', level=1)

            # Add table
            table = doc.add_table(rows=len(desc_stats_df) + 1, cols=len(desc_stats_df.columns))
            table.style = 'Light Grid Accent 1'

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(desc_stats_df.columns):
                hdr_cells[i].text = col_name
                # Make header bold
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Data rows
            for i, row in desc_stats_df.iterrows():
                row_cells = table.rows[i + 1].cells
                for j, val in enumerate(row.values):
                    row_cells[j].text = str(val)

            # Add note
            note = doc.add_paragraph()
            note.add_run('Note: ').bold = True
            note.add_run('CI = confidence interval')

            # Section 2: Organization Summary (if available)
            if org_summary_df is not None and not org_summary_df.empty:
                doc.add_page_break()
                doc.add_heading('Organization Summary', level=1)

                # Add table
                table2 = doc.add_table(rows=len(org_summary_df) + 1, cols=len(org_summary_df.columns))
                table2.style = 'Light Grid Accent 1'

                # Header row
                hdr_cells = table2.rows[0].cells
                for i, col_name in enumerate(org_summary_df.columns):
                    hdr_cells[i].text = col_name
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Data rows
                for i, row in org_summary_df.iterrows():
                    row_cells = table2.rows[i + 1].cells
                    for j, val in enumerate(row.values):
                        row_cells[j].text = str(val)

            return doc

        # SECTION 1: DESCRIPTIVE STATISTICS
        st.subheader("Section 1: Descriptive Statistics")

        # Identify columns for statistics
        stats_columns = []

        # Add sentiment column if mapped
        if st.session_state.sentiment_column:
            stats_columns.append(st.session_state.sentiment_column)

        # Look for divergence column
        for col in df.columns:
            if 'divergence' in col.lower():
                stats_columns.append(col)
                break

        # Look for emotion columns
        emotion_names = ['joy', 'sadness', 'anger', 'fear', 'surprise', 'disgust']
        for emotion in emotion_names:
            for col in df.columns:
                if emotion in col.lower():
                    stats_columns.append(col)
                    break

        if stats_columns:
            desc_stats_df = calculate_descriptive_stats(df, stats_columns)

            # Style the dataframe with alternating row colors
            def style_dataframe(df):
                styles = []
                for i in range(len(df)):
                    if i % 2 == 0:
                        styles.append(['background-color: #f0f2f6'] * len(df.columns))
                    else:
                        styles.append(['background-color: white'] * len(df.columns))
                return styles

            # Display the table
            st.dataframe(
                desc_stats_df,
                use_container_width=True,
                hide_index=True
            )

            st.caption("*Note: CI = confidence interval*")

            # Store for export
            st.session_state.desc_stats_df = desc_stats_df
        else:
            st.info("No statistical columns found. Please ensure your data contains sentiment scores and/or emotion columns.")
            st.session_state.desc_stats_df = None

        # SECTION 2: ORGANIZATION SUMMARY
        st.divider()
        st.subheader("Section 2: Organization Summary")

        if st.session_state.org_column and st.session_state.sentiment_column:
            org_summary_df = calculate_org_summary(
                df,
                st.session_state.org_column,
                st.session_state.sentiment_column
            )

            # Display the table
            st.dataframe(
                org_summary_df,
                use_container_width=True,
                hide_index=True
            )

            # Store for export
            st.session_state.org_summary_df = org_summary_df
        else:
            st.info("Organization and sentiment columns must be mapped to display this summary. Please go to the 'Column Mapping' tab.")
            st.session_state.org_summary_df = None

        # SECTION 3: DOWNLOAD BUTTONS
        st.divider()
        st.subheader("Section 3: Download Options")

        col1, col2, col3, col4, col5 = st.columns(5)

        # Download Results CSV
        with col1:
            csv_buffer = BytesIO()
            df.to_csv(csv_buffer, index=False)
            csv_buffer.seek(0)

            st.download_button(
                label="📥 Download Results CSV",
                data=csv_buffer,
                file_name="sentiment_results.csv",
                mime="text/csv",
                use_container_width=True,
                help="Download the complete results dataframe"
            )

        # Download LaTeX Tables
        with col2:
            if st.session_state.get('desc_stats_df') is not None:
                latex_content = "% LaTeX Tables for Sentiment Analysis\n"
                latex_content += "% Requires \\usepackage{booktabs}\n\n"

                # Add descriptive statistics table
                latex_content += create_latex_table(
                    st.session_state.desc_stats_df,
                    "Descriptive Statistics",
                    "tab:descriptive_stats"
                )
                latex_content += "\n\n"

                # Add organization summary table if available
                if st.session_state.get('org_summary_df') is not None:
                    latex_content += create_latex_table(
                        st.session_state.org_summary_df,
                        "Organization Summary",
                        "tab:org_summary"
                    )

                st.download_button(
                    label="📄 Download Tables (LaTeX)",
                    data=latex_content,
                    file_name="tables.tex",
                    mime="text/plain",
                    use_container_width=True,
                    help="Download formatted LaTeX tables"
                )
            else:
                st.button(
                    label="📄 Download Tables (LaTeX)",
                    disabled=True,
                    use_container_width=True,
                    help="No statistics available"
                )

        # Download Word Tables
        with col3:
            if st.session_state.get('desc_stats_df') is not None:
                word_buffer = BytesIO()
                doc = create_word_document(
                    st.session_state.desc_stats_df,
                    st.session_state.get('org_summary_df')
                )
                doc.save(word_buffer)
                word_buffer.seek(0)

                st.download_button(
                    label="📝 Download Tables (Word)",
                    data=word_buffer,
                    file_name="tables.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="Download formatted Word document with tables"
                )
            else:
                st.button(
                    label="📝 Download Tables (Word)",
                    disabled=True,
                    use_container_width=True,
                    help="No statistics available"
                )

        # Download All Figures (ZIP)
        with col4:
            if st.button("🖼️ Download All Figures (ZIP)", use_container_width=True, help="Download all visualizations as PDFs"):
                with st.spinner("Creating ZIP file..."):
                    try:
                        # Create temporary directory
                        with tempfile.TemporaryDirectory() as temp_dir:
                            zip_path = os.path.join(temp_dir, "figures.zip")

                            with zipfile.ZipFile(zip_path, 'w') as zipf:
                                # Check for existing exports
                                exports_dir = 'exports'
                                if os.path.exists(exports_dir):
                                    for file in os.listdir(exports_dir):
                                        if file.endswith(('.pdf', '.png')):
                                            file_path = os.path.join(exports_dir, file)
                                            zipf.write(file_path, arcname=file)

                                # If heatmap data exists, generate and add it
                                if st.session_state.get('heatmap_data') is not None:
                                    # Generate heatmap if not already exported
                                    heatmap_path = os.path.join(exports_dir, 'sentiment_heatmap.pdf')
                                    if not os.path.exists(heatmap_path):
                                        from app import create_publication_heatmap
                                        fig = create_publication_heatmap(
                                            st.session_state.heatmap_data,
                                            st.session_state.get('row_linkage'),
                                            st.session_state.get('col_linkage'),
                                            (10, 7),
                                            300
                                        )
                                        fig.savefig(heatmap_path, format='pdf', bbox_inches='tight', dpi=300)
                                        plt.close(fig)
                                        zipf.write(heatmap_path, arcname='sentiment_heatmap.pdf')

                            # Read the zip file
                            with open(zip_path, 'rb') as f:
                                zip_data = f.read()

                            st.download_button(
                                label="⬇️ Download ZIP",
                                data=zip_data,
                                file_name="all_figures.zip",
                                mime="application/zip",
                                use_container_width=True
                            )
                            st.success("✅ ZIP file created successfully!")
                    except Exception as e:
                        st.error(f"❌ Error creating ZIP file: {str(e)}")

        # Download Replication Package (ZIP)
        with col5:
            if st.button("📦 Download Replication Package", use_container_width=True, help="Download data + figures + methodology notes"):
                with st.spinner("Creating replication package..."):
                    try:
                        # Create temporary directory
                        with tempfile.TemporaryDirectory() as temp_dir:
                            zip_path = os.path.join(temp_dir, "replication_package.zip")

                            with zipfile.ZipFile(zip_path, 'w') as zipf:
                                # Add data
                                data_path = os.path.join(temp_dir, 'data.csv')
                                df.to_csv(data_path, index=False)
                                zipf.write(data_path, arcname='data/results.csv')

                                # Add statistics tables
                                if st.session_state.get('desc_stats_df') is not None:
                                    stats_path = os.path.join(temp_dir, 'descriptive_statistics.csv')
                                    st.session_state.desc_stats_df.to_csv(stats_path, index=False)
                                    zipf.write(stats_path, arcname='tables/descriptive_statistics.csv')

                                if st.session_state.get('org_summary_df') is not None:
                                    org_path = os.path.join(temp_dir, 'organization_summary.csv')
                                    st.session_state.org_summary_df.to_csv(org_path, index=False)
                                    zipf.write(org_path, arcname='tables/organization_summary.csv')

                                # Add figures from exports directory
                                exports_dir = 'exports'
                                if os.path.exists(exports_dir):
                                    for file in os.listdir(exports_dir):
                                        if file.endswith(('.pdf', '.png')):
                                            file_path = os.path.join(exports_dir, file)
                                            zipf.write(file_path, arcname=f'figures/{file}')

                                # Add methodology notes
                                methodology = """SENTIMENT ANALYSIS REPLICATION PACKAGE
========================================

Contents:
---------
1. data/results.csv - Complete sentiment analysis results
2. tables/ - Statistical tables (descriptive statistics and organization summaries)
3. figures/ - All visualizations in PDF/PNG format

Methodology Notes:
------------------
- Sentiment scores range from -1 (negative) to +1 (positive)
- Divergence scores measure the mismatch between sentiment and emotional content
- Emotion scores are normalized to range [0, 1]
- 95% confidence intervals calculated using t-distribution
- Hierarchical clustering performed using Ward's method with Euclidean distance

Column Mapping:
---------------
"""
                                if st.session_state.org_column:
                                    methodology += f"- Organization: {st.session_state.org_column}\n"
                                if st.session_state.question_column:
                                    methodology += f"- Question: {st.session_state.question_column}\n"
                                if st.session_state.sentiment_column:
                                    methodology += f"- Sentiment: {st.session_state.sentiment_column}\n"

                                methodology += "\nGenerated by Sentiment Analysis Dashboard\n"

                                methodology_path = os.path.join(temp_dir, 'README.txt')
                                with open(methodology_path, 'w') as f:
                                    f.write(methodology)
                                zipf.write(methodology_path, arcname='README.txt')

                            # Read the zip file
                            with open(zip_path, 'rb') as f:
                                zip_data = f.read()

                            st.download_button(
                                label="⬇️ Download Package",
                                data=zip_data,
                                file_name="replication_package.zip",
                                mime="application/zip",
                                use_container_width=True
                            )
                            st.success("✅ Replication package created successfully!")
                    except Exception as e:
                        st.error(f"❌ Error creating replication package: {str(e)}")

# Footer
st.divider()
st.markdown(
    """
    <div style='text-align: center; color: gray; padding: 20px;'>
        <p>Sentiment Analysis Dashboard | Built with Streamlit</p>
    </div>
    """,
    unsafe_allow_html=True
)
